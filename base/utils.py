from docx import Document
from .models import Question, Question2
from docx import Document



def process_question_file(file_path):
    document = Document(file_path)
    questions = []
    question_type = None  # Loại câu hỏi: 1 - Trắc nghiệm, 2 - Đúng/Sai
    question = None

    if not document.paragraphs or all(not p.text.strip() for p in document.paragraphs):
        raise ValueError("File không có nội dung hoặc không đúng định dạng.")

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        # Xác định phần loại câu hỏi
        if text.startswith("PHẦN 1: TRẮC NGHIỆM"):
            question_type = 1
            continue
        elif text.startswith("PHẦN 2: CÂU TRẮC NGHIỆM ĐÚNG/SAI"):
            question_type = 2
            continue

        # Phát hiện câu hỏi
        if text.startswith("Câu "):
            if question:  # Lưu câu hỏi trước đó
                questions.append(question)
            question = {
                "type": question_type,
                "name": text.split(": ", 1)[1].strip() if ": " in text else "",
                "Ans_a": "",
                "Ans_b": "",
                "Ans_c": "",
                "Ans_d": "",
                "correct": None,
                "correct_a": "False",
                "correct_b": "False",
                "correct_c": "False",
                "correct_d": "False",
            }

        # Phát hiện đáp án trắc nghiệm
        elif text.startswith(("A.", "B.", "C.", "D.")) and question_type == 1:
            key = f"Ans_{text[0].lower()}"
            answer_text = text[2:].strip()
            if answer_text:
                is_correct = any(run.bold for run in paragraph.runs)
                question[key] = answer_text
                if is_correct:
                    question["correct"] = key

        # Phát hiện đáp án đúng/sai
        elif text.startswith(("A.", "B.", "C.", "D.")) and question_type == 2:
            key = f"correct_{text[0].lower()}"
            answer_text = text[2:].strip()
            if answer_text:
                is_correct = any(run.bold for run in paragraph.runs)
                question[f"Ans_{text[0].lower()}"] = answer_text
                question[key] = "True" if is_correct else "False"

    if question:
        questions.append(question)  # Lưu câu hỏi cuối cùng

    return questions
def save_questions_to_db(questions, de, loai):
    for q in questions:
        if q["type"] == 1:  # Câu hỏi trắc nghiệm
            Question2.objects.create(
                Noi_dung=q["name"],
                A=q["Ans_a"],
                B=q["Ans_b"],
                C=q["Ans_c"],
                D=q["Ans_d"],
                Corect_ans_single=q["correct"],
                De=de,
                Loai=loai,
                type=1
            )
        elif q["type"] == 2:  # Câu hỏi đúng/sai
            Question2.objects.create(
                Noi_dung=q["name"],
                A=q["Ans_a"],
                B=q["Ans_b"],
                C=q["Ans_c"],
                D=q["Ans_d"],
                Corect_ans_a=q["correct_a"],
                Corect_ans_b=q["correct_b"],
                Corect_ans_c=q["correct_c"],
                Corect_ans_d=q["correct_d"],
                De=de,
                Loai=loai,
                type=2
            )